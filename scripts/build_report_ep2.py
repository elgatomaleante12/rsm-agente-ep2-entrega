from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs"
REPORT_DOCX = DOCS_DIR / "Informe_EP2_RSM_Agente_Funcional.docx"
DIAGRAM_PNG = DOCS_DIR / "diagrama_orquestacion.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 32, 42)
MUTED = RGBColor(95, 105, 120)
LIGHT_FILL = "F2F4F7"
PALE_BLUE = "E8EEF5"


def set_run_font(run, size: float = 11, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.space_after = Pt(8)
    list_bullet.paragraph_format.line_spacing = 1.167


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def add_table(doc: Document, header: list[str], rows_: list[list[str]], widths_dxa: list[int], fill: str = LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        hdr[i].text = text
        set_cell_shading(hdr[i], fill)
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, bold=True, color=INK)
    for row in rows_:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()
    return table


def add_paragraph(doc: Document, text: str, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        run = p.add_run(bold_label)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_label) :])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def draw_diagram(path: Path) -> None:
    width, height = 1400, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 30)
        font = ImageFont.truetype("arial.ttf", 25)
        small = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        title_font = ImageFont.load_default()
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    def box(x1, y1, x2, y2, label, fill="#F2F4F7", outline="#506176", title=False):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=outline, width=3)
        lines = label.split("\n")
        total_h = len(lines) * 30
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=title_font if title else font)
            draw.text(((x1 + x2 - bbox[2]) / 2, y1 + (y2 - y1 - total_h) / 2 + idx * 32), line, fill="#17202A", font=title_font if title else font)

    def arrow(start, end, color="#0F6B5F"):
        draw.line((start, end), fill=color, width=5)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) > abs(ey - sy):
            direction = 1 if ex > sx else -1
            pts = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
        else:
            direction = 1 if ey > sy else -1
            pts = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
        draw.polygon(pts, fill=color)

    draw.text((40, 30), "Orquestación del agente RSM EP2", fill="#1F4D78", font=title_font)
    box(60, 120, 285, 220, "Usuario\nmecánico", "#E8EEF5")
    box(385, 120, 640, 220, "FastAPI\nWeb UI", "#F2F4F7")
    box(740, 120, 1000, 220, "RSMAgent\ncontrolador", "#E8EEF5")
    box(1085, 120, 1330, 220, "TaskPlanner\ndecisión", "#F2F4F7")

    box(95, 360, 330, 470, "Consulta\ninventario", "#FFFFFF")
    box(385, 360, 620, 470, "Escritura\nsalida / reporte", "#FFFFFF")
    box(675, 360, 910, 470, "Razonamiento\nalertas", "#FFFFFF")
    box(965, 360, 1200, 470, "Contexto\nsemántico", "#FFFFFF")

    box(155, 610, 385, 700, "SQLite\nstock + OT", "#F7FBFA")
    box(525, 610, 755, 700, "Memoria\ncorto/largo plazo", "#F7FBFA")
    box(895, 610, 1125, 700, "Reglas\noperativas", "#F7FBFA")

    arrow((285, 170), (385, 170))
    arrow((640, 170), (740, 170))
    arrow((1000, 170), (1085, 170))
    arrow((1205, 220), (1085, 360))
    arrow((1090, 220), (795, 360))
    arrow((865, 220), (505, 360))
    arrow((770, 220), (215, 360))
    arrow((215, 470), (270, 610))
    arrow((505, 470), (270, 610))
    arrow((505, 470), (640, 610))
    arrow((795, 470), (640, 610))
    arrow((1080, 470), (1010, 610))
    arrow((1080, 470), (640, 610))
    draw.text((55, 725), "Evidencia: el agente devuelve plan, herramientas usadas, memoria reciente y respuesta con fuente.", fill="#5F6978", font=small)
    image.save(path)


def build_docx() -> Path:
    DOCS_DIR.mkdir(exist_ok=True)
    draw_diagram(DIAGRAM_PNG)

    doc = Document()
    configure_styles(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("RSM EP2 | Informe técnico")
    set_run_font(run, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("RSM - Agente Funcional de Gestión de Inventario")
    set_run_font(run, size=22, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Informe técnico EP2 | ISY0101 | Versión funcional en Visual Studio Code")
    set_run_font(run, size=12.5, color=MUTED)

    add_table(
        doc,
        ["Dato", "Detalle"],
        [
            ["Proyecto", "Sistema inteligente RSM para consultas, trazabilidad y alertas de stock."],
            ["Repositorio local", "Carpeta rsm-agente-ep2, lista para abrirse en VS Code o subirse a GitHub/GitLab."],
            ["Stack", "Python 3, FastAPI, SQLite, LangChain Core, HTML/CSS/JS y pytest."],
        ],
        [2000, 7360],
        fill=PALE_BLUE,
    )

    doc.add_heading("1. Propósito y alcance", level=1)
    add_paragraph(
        doc,
        "El proyecto transforma el diseño de EP1 en un MVP funcional para el taller RSM. "
        "El agente permite consultar inventario en lenguaje natural, registrar salidas con orden de trabajo, "
        "detectar quiebres de stock y redactar evidencia operativa. La solucion opera localmente para respetar "
        "la restricción de privacidad del caso y evitar dependencia de APIs pagadas."
    )
    add_paragraph(
        doc,
        "Flujo principal: el mecánico envía una solicitud, FastAPI la entrega al agente, el planificador decide "
        "el flujo y las herramientas LangChain consultan o escriben datos en SQLite. La respuesta incluye plan, "
        "herramientas utilizadas y memoria reciente, por lo que el evaluador puede validar la orquestación."
    )
    doc.add_picture(str(DIAGRAM_PNG), width=Inches(6.25))

    add_page_break(doc)
    doc.add_heading("2. Diseño e implementación del agente", level=1)
    add_paragraph(
        doc,
        "La implementacion se separa en cinco modulos: API, agente, planificador, herramientas y persistencia. "
        "Esta separación permite probar cada responsabilidad y demostrar autonomía sin ocultar la lógica en un script único."
    )
    add_table(
        doc,
        ["Componente", "Rol técnico", "Evidencia"],
        [
            ["app/main.py", "Expone API REST, interfaz web y endpoints de salud, inventario, chat, memoria y reportes.", "FastAPI + /docs"],
            ["app/agent.py", "Orquesta plan, herramientas y memoria para cada mensaje del usuario.", "Respuesta con plan y tools_used"],
            ["app/tools.py", "Define herramientas LangChain de consulta, escritura, contexto y alertas.", "StructuredTool"],
            ["app/database.py", "Gestiona SQLite, inventario, movimientos y memoria persistente.", "data/rsm_agent.db"],
            ["app/retriever.py", "Recupera contexto local desde inventario, reglas y memoria por similitud.", "hits con score"],
        ],
        [1900, 4650, 2810],
    )
    doc.add_heading("Herramientas integradas", level=2)
    add_bullet(doc, "Consulta: consultar_inventario recupera repuestos, stock, ubicación, proveedor y compatibilidad.")
    add_bullet(doc, "Escritura: registrar_salida descuenta stock con trazabilidad y redactar_reporte_alertas genera evidencia en outputs.")
    add_bullet(doc, "Razonamiento: generar_alertas clasifica CRÍTICO o BAJO y recomienda reposición según prioridad.")

    add_page_break(doc)
    doc.add_heading("3. Memoria, contexto y planificación", level=1)
    add_paragraph(
        doc,
        "La memoria corta se representa con los últimos eventos de la sesión; la memoria larga queda persistida en la tabla "
        "memory de SQLite. Cada salida registrada guarda una huella con OT, mecánico, vehículo, repuesto y stock resultante. "
        "La recuperación de contexto combina tres fuentes: inventario vigente, reglas operativas y memoria histórica de la sesión."
    )
    add_table(
        doc,
        ["Escenario", "Decisión adaptativa del agente", "Resultado esperado"],
        [
            ["Consulta de stock", "Busca contexto relevante y responde con fuente.", "Stock, ubicación, proveedor y compatibilidad."],
            ["Salida incompleta", "Detecta campos faltantes antes de escribir.", "No descuenta stock y solicita OT/mecánico/vehículo."],
            ["Stock insuficiente", "Valida disponibilidad antes del movimiento.", "Rechaza la salida y conserva inventario."],
            ["Stock bajo mínimo", "Registra salida y agrega advertencia operacional.", "Recomienda reposición al proveedor."],
        ],
        [2100, 4200, 3060],
    )
    doc.add_heading("Planificación de tareas", level=2)
    add_paragraph(
        doc,
        "El módulo app/planner.py clasifica la intención en consultar_stock, registrar_salida, generar_alertas, "
        "recomendar_reposicion o redactar_reporte. Para cada intención secuencia pasos: validar datos, consultar contexto, "
        "ejecutar herramienta, registrar memoria y responder con evidencia."
    )

    add_page_break(doc)
    doc.add_heading("4. Evidencia de funcionamiento", level=1)
    add_paragraph(
        doc,
        "El proyecto fue preparado para Visual Studio Code con archivo .code-workspace, launch.json y tasks.json. "
        "El evaluador puede instalar dependencias, iniciar la API y ejecutar pruebas desde Terminal > Run Task o con F5."
    )
    add_table(
        doc,
        ["Prueba", "Comando o entrada", "Evidencia"],
        [
            ["Tests automatizados", ".\\.venv\\Scripts\\python.exe -m pytest -q", "6 passed"],
            ["Consulta", "¿Cuántos filtros de aceite Toyota quedan?", "Usa recuperar_contexto y consultar_inventario."],
            ["Salida valida", "Salida de 2 unidades FO-TOY-001 para OT-778", "Stock baja de 8 a 6 y se guarda memoria."],
            ["Salida sin trazabilidad", "Registra salida de 2 unidades FO-TOY-001", "Bloquea escritura por falta de OT/mecánico/vehículo."],
            ["Alertas", "¿Qué ítems están críticos?", "Entrega ítems CRÍTICO/BAJO y proveedor sugerido."],
        ],
        [2100, 3950, 3310],
    )
    doc.add_heading("Instrucciones de ejecución", level=2)
    add_bullet(doc, "Abrir rsm-agente-ep2.code-workspace en Visual Studio Code.")
    add_bullet(doc, "Ejecutar las tareas: Instalar dependencias, Inicializar base de datos y Ejecutar API.")
    add_bullet(doc, "Probar la interfaz en http://127.0.0.1:8000 o en http://<IP_DEL_EQUIPO>:8000 y la documentación en /docs.")

    add_page_break(doc)
    doc.add_heading("5. Cierre y referencias", level=1)
    add_paragraph(
        doc,
        "La versión EP2 cumple el paso clave que faltaba en el diseño inicial: existe un agente ejecutable, verificable "
        "y con persistencia real. El sistema no simula solo una respuesta de IA; aplica reglas, consulta datos, escribe "
        "movimientos y conserva memoria. La solución puede crecer incorporando un LLM local por Ollama, ChromaDB o un frontend "
        "React, pero el MVP actual ya permite validar autonomía funcional y continuidad de tarea."
    )
    add_table(
        doc,
        ["Indicador", "Cumplimiento"],
        [
            ["IE1-IE2", "Herramientas autónomas implementadas con LangChain Core."],
            ["IE3-IE4", "Memoria persistente y recuperación de contexto local."],
            ["IE5-IE6", "Planificación por intención y decisiones ante condiciones cambiantes."],
            ["IE7-IE10", "README, diagrama, pruebas, informe técnico y referencias."],
        ],
        [1800, 7560],
    )
    doc.add_heading("Referencias", level=2)
    refs = [
        "FastAPI. (2024). FastAPI documentation. https://fastapi.tiangolo.com/",
        "LangChain. (2024). LangChain documentation. https://python.langchain.com/docs/",
        "Python Software Foundation. (2024). Python documentation. https://docs.python.org/",
        "SQLite. (2024). SQLite documentation. https://sqlite.org/docs.html",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    doc.save(REPORT_DOCX)
    return REPORT_DOCX


if __name__ == "__main__":
    print(build_docx())
