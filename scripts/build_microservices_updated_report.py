from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs"
OUT_DOCX = DOCS_DIR / "RSM_Taller_final_10_microservicios_actualizado_informe.docx"
MICRO_IMG = DOCS_DIR / "mapa_microservicios_actualizado.png"
AGENT_IMG = DOCS_DIR / "diagrama_orquestacion.png"

GITHUB_REPO = "https://github.com/elgatomaleante12/rsm-agente-ep2-entrega"

BLUE = RGBColor(32, 103, 168)
DARK = RGBColor(20, 28, 38)
MUTED = RGBColor(92, 103, 118)
FILL = "EEF3F8"


def set_run(run, size=11, bold=False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(0.82))
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Heading 1", 16, 14, 7),
        ("Heading 2", 13, 10, 5),
        ("Heading 3", 11.5, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if name != "Heading 3" else DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("RSM | Informe actualizado")
    set_run(run, size=9, color=MUTED)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)


def table(doc: Document, header: list[str], data: list[list[str]], widths: list[int], fill=FILL):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = text
        shade(c, fill)
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run(r, bold=True, color=DARK)
    for row in data:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run(r)
    set_table_widths(t, widths)
    doc.add_paragraph()
    return t


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r)


def make_micro_map(path: Path) -> None:
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 24)
        small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        title_font = ImageFont.load_default()
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    draw.text((50, 35), "Mapa de servicios RSM", fill="#2067A8", font=title_font)
    services = [
        ("cliente-service", "8081", 60, 135),
        ("vehiculo-service", "8082", 365, 135),
        ("ordendetrabajo", "8083", 670, 135),
        ("mecanico-service", "8084", 975, 135),
        ("asignacion-service", "8085", 1280, 135),
        ("repuesto-service", "8086", 60, 395),
        ("detalleorden-service", "8087", 365, 395),
        ("factura-service", "8088", 670, 395),
        ("pago-service", "8089", 975, 395),
        ("historial-service", "8090", 1280, 395),
    ]

    for name, port, x, y in services:
        draw.rounded_rectangle((x, y, x + 190, y + 115), radius=18, fill="#EEF3F8", outline="#4E6072", width=3)
        draw.text((x + 18, y + 24), name, fill="#141C26", font=small)
        draw.text((x + 70, y + 66), port, fill="#0F6B5F", font=font)

    def arrow(x1, y1, x2, y2):
        draw.line((x1, y1, x2, y2), fill="#0F6B5F", width=5)
        draw.polygon([(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)], fill="#0F6B5F")

    arrow(250, 190, 365, 190)
    arrow(555, 190, 670, 190)
    arrow(860, 190, 975, 190)
    arrow(1165, 190, 1280, 190)
    arrow(250, 450, 365, 450)
    arrow(555, 450, 670, 450)
    arrow(860, 450, 975, 450)
    arrow(1165, 450, 1280, 450)
    draw.rounded_rectangle((255, 650, 1245, 790), radius=20, fill="#F7FBFA", outline="#0F6B5F", width=3)
    draw.text((295, 685), "Flujo: cliente -> vehículo -> orden -> mecánico -> repuesto -> detalle -> factura -> pago -> historial", fill="#141C26", font=font)
    draw.text((295, 730), "Cada servicio mantiene su dominio y se comunica por IDs referenciados.", fill="#5C6776", font=small)
    img.save(path)


def build() -> Path:
    DOCS_DIR.mkdir(exist_ok=True)
    make_micro_map(MICRO_IMG)

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Sistema de Gestión de Inventario RSM")
    set_run(r, size=24, bold=True, color=DARK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    r = sub.add_run("Informe actualizado: microservicios + agente funcional de IA")
    set_run(r, size=14, color=MUTED)

    table(
        doc,
        ["Campo", "Información actualizada"],
        [
            ["Asignatura", "INGENIERÍA DE SOLUCIONES CON INTELIGENCIA ARTIFICIAL_001D"],
            ["Evaluación", "Evaluación 2"],
            ["Docente", "HECTOR ANDRES MOREL BRIONES"],
            ["Nombre de equipo", "Joaquin Alvarado"],
            ["Repositorio GitHub", GITHUB_REPO],
        ],
        [2500, 7400],
        fill="E6EEF7",
    )

    doc.add_heading("1. Problemática y contexto", level=1)
    para(
        doc,
        "Los talleres mecánicos pequeños suelen operar con cuadernos, carpetas físicas y registros en papel. "
        "Esto genera fallos frecuentes en el control de stock, errores en pedidos de repuestos, baja trazabilidad "
        "de órdenes de trabajo y poca transparencia para el cliente sobre el estado de su vehículo."
    )
    para(
        doc,
        "El impacto operativo se observa en retrasos de reparación, duplicación de tareas, pérdida de seguimiento "
        "y tiempos muertos al buscar disponibilidad de repuestos. El sistema RSM responde a este problema mediante "
        "digitalización de procesos, arquitectura de microservicios y un agente funcional de IA para consultar, razonar "
        "y registrar información del inventario."
    )

    doc.add_heading("2. Visión del proyecto", level=1)
    para(
        doc,
        "RSM busca transformar la gestión de talleres mecánicos mediante la digitalización completa del flujo operativo. "
        "La solución considera sincronización en tiempo real, historial por patente, estados actualizados y trazabilidad "
        "desde el ingreso del cliente hasta el cierre de la orden."
    )
    bullet(doc, "Sincronización de datos entre servicios mediante IDs referenciados.")
    bullet(doc, "Historial técnico por vehículo para mantener continuidad operacional.")
    bullet(doc, "Control de repuestos, costos, facturación y pagos en servicios separados.")
    bullet(doc, "Agente de IA para consultas de stock, alertas, reportes y registro de salidas.")

    doc.add_heading("3. Arquitectura de microservicios", level=1)
    para(
        doc,
        "El sistema implementa 10 microservicios independientes. Cada servicio cubre una parte del flujo del taller y "
        "puede evolucionar de forma modular sin afectar directamente a los demás componentes."
    )
    doc.add_picture(str(MICRO_IMG), width=Inches(6.9))
    table(
        doc,
        ["Puerto", "Microservicio", "Responsabilidad"],
        [
            ["8081", "cliente-service", "Registra clientes y datos de contacto."],
            ["8082", "vehiculo-service", "Asocia vehículos con clientes mediante clienteId."],
            ["8083", "ordendetrabajo-service", "Crea órdenes con vehículo, mecánico, estado y descripción."],
            ["8084", "mecanico-service", "Administra mecánicos, especialidad y contacto."],
            ["8085", "asignacion-service", "Asigna mecánicos a vehículos o trabajos del taller."],
            ["8086", "repuesto-service", "Controla repuestos, precios y stock disponible."],
            ["8087", "detalleorden-service", "Registra costos, cantidades y detalles de cada orden."],
            ["8088", "factura-service", "Genera facturas vinculadas a una orden de trabajo."],
            ["8089", "pago-service", "Registra pagos asociados a facturas."],
            ["8090", "historial-service", "Guarda historial técnico de cada vehículo."],
        ],
        [1000, 2600, 6300],
    )

    doc.add_heading("4. Modelado estructural y flujo de negocio", level=1)
    para(
        doc,
        "Cada microservicio mantiene su propio dominio de datos. El modelado usa entidades relacionadas por IDs "
        "referenciados, evitando acoplamiento entre bases de datos y permitiendo despliegues independientes. "
        "Las migraciones de esquema se controlan con Flyway y la persistencia se implementa con Spring Data JPA."
    )
    para(
        doc,
        "El flujo de negocio conecta nueve pasos: registro del cliente, asociación del vehículo, creación de orden, "
        "asignación de mecánico, gestión de repuestos, cálculo de detalle y costos, generación de factura, registro "
        "de pago y actualización del historial del vehículo."
    )
    table(
        doc,
        ["Capa", "Tecnología", "Rol"],
        [
            ["Backend microservicios", "Java 21 + Spring Boot 3.x", "APIs REST para cada dominio del taller."],
            ["Persistencia", "Spring Data JPA + Hibernate ORM", "Mapeo objeto-relacional y repositorios."],
            ["Base de datos", "MariaDB puerto 5506", "Almacenamiento transaccional por servicio."],
            ["Migraciones", "Flyway", "Versionado automático de esquemas."],
            ["Validación", "Bean Validation", "Control de reglas y datos de entrada."],
            ["Construcción", "Maven", "Gestión de dependencias y ciclo de build."],
        ],
        [2200, 3200, 4500],
    )

    doc.add_heading("5. Agente funcional de IA agregado", level=1)
    para(
        doc,
        "La versión actual incorpora un agente funcional ejecutable en Visual Studio Code. Este agente no reemplaza "
        "la arquitectura de microservicios; la complementa como capa inteligente para consulta, escritura, recuperación "
        "de contexto y toma de decisiones sobre inventario."
    )
    table(
        doc,
        ["Dato", "Detalle"],
        [
            ["Proyecto", "RSM - Agente Funcional EP2 para inventario mecánico."],
            ["Repositorio GitHub", GITHUB_REPO],
            ["Ejecución local", "Abrir rsm-agente-ep2.code-workspace y ejecutar las tareas de VS Code."],
            ["API", "FastAPI disponible en http://127.0.0.1:8000 y documentación en /docs."],
            ["Pruebas", "pytest validado con resultado: 6 passed."],
        ],
        [2200, 7700],
        fill="E6EEF7",
    )
    doc.add_picture(str(AGENT_IMG), width=Inches(6.8))

    doc.add_heading("6. Herramientas, memoria y planificación del agente", level=1)
    para(
        doc,
        "El agente se implementa en Python con FastAPI, SQLite y LangChain Core. La clase RSMAgent orquesta un "
        "planificador determinístico y cinco herramientas: consultar inventario, registrar salida, generar alertas, "
        "recuperar contexto y redactar reportes operativos."
    )
    bullet(doc, "Consulta: recupera stock, ubicación, proveedor y compatibilidad de repuestos.")
    bullet(doc, "Escritura: registra salidas de inventario con orden de trabajo, mecánico y vehículo.")
    bullet(doc, "Razonamiento: clasifica stock CRÍTICO o BAJO y sugiere reposición.")
    bullet(doc, "Memoria: guarda eventos de usuario, respuestas, movimientos y reportes en SQLite.")
    bullet(doc, "Recuperación de contexto: combina inventario, reglas operativas y memoria previa.")
    para(
        doc,
        "La planificación ajusta el comportamiento ante condiciones cambiantes. Si faltan datos de trazabilidad, "
        "el agente no descuenta stock; si la cantidad solicitada supera la disponibilidad, rechaza el movimiento; "
        "si el stock queda bajo mínimo, registra el evento y advierte la necesidad de reposición."
    )

    doc.add_heading("7. Evidencia de funcionamiento", level=1)
    table(
        doc,
        ["Escenario", "Entrada", "Resultado"],
        [
            ["Consulta de stock", "¿Cuántos filtros de aceite Toyota quedan?", "Responde FO-TOY-001 con stock, ubicación y proveedor."],
            ["Alertas", "¿Qué ítems están críticos o bajo mínimo?", "Lista ítems CRÍTICO/BAJO con proveedor sugerido."],
            ["Salida válida", "Salida de 2 unidades FO-TOY-001 para OT-778", "Descuenta stock y guarda memoria de trazabilidad."],
            ["Salida incompleta", "Registra salida de 2 unidades FO-TOY-001", "Bloquea escritura por falta de OT, mecánico o vehículo."],
            ["Reporte", "Genera un reporte de alertas de stock", "Crea archivo Markdown en outputs/."],
        ],
        [1900, 3950, 4050],
    )

    doc.add_heading("8. Conclusión", level=1)
    para(
        doc,
        "El sistema RSM queda documentado como una solución integral: por un lado, presenta una arquitectura fullstack "
        "basada en microservicios para digitalizar el flujo completo del taller; por otro, incorpora un agente de IA "
        "funcional que permite consultar inventario, tomar decisiones operativas, registrar movimientos y conservar "
        "memoria de contexto. Esta combinación responde al requerimiento de construir una solución real, ejecutable, "
        "verificable y alineada con Ingeniería de Soluciones con Inteligencia Artificial."
    )

    doc.add_heading("Referencias", level=1)
    for ref in [
        "FastAPI. (2024). FastAPI documentation. https://fastapi.tiangolo.com/",
        "LangChain. (2024). LangChain documentation. https://python.langchain.com/docs/",
        "Oracle. (2024). Java documentation. https://docs.oracle.com/en/java/",
        "Spring. (2024). Spring Boot documentation. https://spring.io/projects/spring-boot",
        "SQLite. (2024). SQLite documentation. https://sqlite.org/docs.html",
    ]:
        para(doc, ref)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
